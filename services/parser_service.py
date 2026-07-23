import os
import re
import json
import logging
from typing import Dict, Any, Optional
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    from google import genai
except ImportError:
    genai = None

try:
    from openai import OpenAI
except ImportError:
    OpenAI = None

try:
    from anthropic import Anthropic
except ImportError:
    Anthropic = None

import io
from typing import Dict, Any, Optional, Union

logger = logging.getLogger("ats-optimizer")

def extract_text_from_pdf(file_or_path: Union[str, io.BytesIO]) -> str:
    """Extracts text from a PDF file using pdfplumber."""
    if not pdfplumber:
        raise Exception("pdfplumber library is not installed in the current Python environment.")
    text = ""
    try:
        if hasattr(file_or_path, "seek"):
            file_or_path.seek(0)
        with pdfplumber.open(file_or_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if hasattr(file_or_path, "seek"):
            file_or_path.seek(0)
    except Exception as e:
        logger.error(f"pdfplumber extraction failed: {e}")
        raise Exception("Failed to extract text from PDF file.")
    return text


def extract_text_from_docx(file_or_path: Union[str, io.BytesIO]) -> str:
    """Extracts text from a DOCX file using python-docx."""
    if not docx:
        raise Exception("python-docx library is not installed in the current Python environment.")
    try:
        if hasattr(file_or_path, "seek"):
            file_or_path.seek(0)
        doc = docx.Document(file_or_path)
        text_runs = []
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text_runs.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text_runs.append(cell.text)
        if hasattr(file_or_path, "seek"):
            file_or_path.seek(0)
        return "\n".join(text_runs)
    except Exception as e:
        logger.error(f"python-docx extraction failed: {e}")
        raise Exception("Failed to extract text from DOCX file.")


def extract_links_from_document(file_or_path: Union[str, io.BytesIO], ext: str) -> list[str]:
    """Extracts embedded hyperlinks from PDF or DOCX without polluting raw text."""
    links = []
    try:
        if hasattr(file_or_path, "seek"):
            file_or_path.seek(0)
            
        if ext == "pdf":
            if not pdfplumber:
                return []
            with pdfplumber.open(file_or_path) as pdf:
                for page in pdf.pages:
                    if hasattr(page, "hyperlinks") and page.hyperlinks:
                        for hl in page.hyperlinks:
                            uri = hl.get("uri")
                            if uri:
                                links.append(uri)
        else: # docx
            if not docx:
                return []
            doc = docx.Document(file_or_path)
            if hasattr(doc, "part") and hasattr(doc.part, "rels"):
                for rel in doc.part.rels.values():
                    if hasattr(rel, "reltype") and "hyperlink" in rel.reltype and hasattr(rel, "target_ref"):
                        links.append(rel.target_ref)
                        
        if hasattr(file_or_path, "seek"):
            file_or_path.seek(0)
    except Exception as e:
        logger.error(f"Error extracting links: {e}")
    return links


PARSING_SYSTEM_INSTRUCTION = """
You are an expert resume parser. Your task is to parse the raw text of a resume into a structured JSON object.
Extract all details accurately without changing any wording, facts, or context. Do not add any new information, and do not use placeholders. If a section is missing, leave it as an empty list or empty string.

Return the response ONLY in a valid JSON object matching this exact schema:
{
  "personalInfo": {
    "name": "Full Name",
    "email": "Email Address",
    "phone": "Phone Number",
    "github": "Github URL (if present, else empty string)",
    "linkedin": "LinkedIn URL (if present, else empty string)",
    "portfolio": "Portfolio or personal website URL (if present, else empty string)"
  },
  "summary": "Professional Summary or Objective (if present, else empty string)",
  "skills": [
    "Exact line 1 from skills section (e.g. 'Programming Languages: Python, SQL' or 'Power BI - DAX, Data Modeling')",
    "Exact line 2 from skills section",
    ...
  ],
  "experience": [
    {
      "role": "Job Title",
      "company": "Company Name",
      "duration": "Dates of employment (e.g., Jun 2021 - Present)",
      "bullets": ["Bullet point 1", "Bullet point 2", ...]
    }
  ],
  "projects": [
    {
      "name": "Project Name",
      "description": "Short description/tech stack (if present, else empty string)",
      "bullets": ["Bullet point 1", "Bullet point 2", ...]
    }
  ],
  "education": [
    {
      "degree": "Degree (e.g., Bachelor of Science in Computer Science)",
      "school": "University/School Name",
      "year": "Graduation Year (e.g., 2023)"
    }
  ]
}
"""

def parse_resume_text_via_llm(raw_text: str) -> Optional[Dict[str, Any]]:
    """Parses raw resume text into structured JSON using LLM APIs."""
    # 1. Try Gemini
    if genai and os.getenv("GEMINI_API_KEY"):
        try:
            logger.info("Using Google Gemini API for parsing...")
            client = genai.Client(api_key=os.getenv("GEMINI_API_KEY"))
            response = client.models.generate_content(
                model='gemini-2.0-flash',
                contents=f"Raw Resume Text:\n{raw_text}",
                config={
                    "response_mime_type": "application/json",
                    "temperature": 0.1,
                    "system_instruction": PARSING_SYSTEM_INSTRUCTION
                }
            )
            return json.loads(response.text or "{}")
        except Exception as e:
            logger.error(f"Gemini parsing failed: {e}")

    # 2. Try OpenAI
    if OpenAI and os.getenv("OPENAI_API_KEY"):
        try:
            logger.info("Using OpenAI GPT API for parsing...")
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response: Any = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": PARSING_SYSTEM_INSTRUCTION},
                    {"role": "user", "content": f"Raw Resume Text:\n{raw_text}"}
                ],
                response_format={"type": "json_object"},
                temperature=0.1,
                stream=False
            )
            return json.loads(response.choices[0].message.content or "{}")
        except Exception as e:
            logger.error(f"OpenAI parsing failed: {e}")

    # 3. Try Claude
    if Anthropic and os.getenv("CLAUDE_API_KEY"):
        try:
            logger.info("Using Anthropic Claude API for parsing...")
            client = Anthropic(api_key=os.getenv("CLAUDE_API_KEY"))
            response: Any = client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=4000,
                system=PARSING_SYSTEM_INSTRUCTION,
                messages=[
                    {"role": "user", "content": f"Raw Resume Text:\n{raw_text}"}
                ],
                temperature=0.1
            )
            content_text = response.content[0].text
            start = content_text.find('{')
            end = content_text.rfind('}') + 1
            if start >= 0 and end > start:
                return json.loads(content_text[start:end])
            return json.loads(content_text)
        except Exception as e:
            logger.error(f"Claude parsing failed: {e}")

    return None


def parse_resume_text_to_structure(text: str, filename: str) -> Dict[str, Any]:
    """Fallback heuristic regex parser for raw resume text."""
    structure: Dict[str, Any] = {
        "personalInfo": {
            "name": filename.replace("_", " ").split(".")[0],
            "email": "",
            "phone": "",
            "github": "",
            "linkedin": "",
            "portfolio": ""
        },
        "summary": "",
        "skills": [],
        "experience": [],
        "projects": [],
        "education": [],
        "certifications": []
    }

    emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    phones = re.findall(r'\+?[0-9\s\-()]{7,18}', text)

    if emails:
        structure["personalInfo"]["email"] = emails[0]
    if phones:
        valid_phones = [p.strip() for p in phones if len(re.sub(r'[^0-9]', '', p)) >= 10]
        if valid_phones:
            structure["personalInfo"]["phone"] = valid_phones[0]

    # Extract LinkedIn URL
    linkedin_match = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_\-\/\.]+', text, re.IGNORECASE)
    if linkedin_match:
        structure["personalInfo"]["linkedin"] = linkedin_match.group(0).strip()

    # Extract GitHub URL
    github_match = re.search(r'(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_\-\/\.]+', text, re.IGNORECASE)
    if github_match:
        structure["personalInfo"]["github"] = github_match.group(0).strip()

    # Extract Portfolio URL (any other valid URL that isn't github/linkedin)
    urls = re.findall(r'(?:https?://)?(?:www\.)?[a-zA-Z0-9.-]+\.[a-zA-Z]{2,6}(?:/[a-zA-Z0-9._%+-]*)*', text)
    for url in urls:
        lower_url = url.lower()
        if "linkedin" not in lower_url and "github" not in lower_url and "email" not in lower_url:
            if not any(lower_url.endswith(ext) for ext in [".pdf", ".docx", ".doc", ".png", ".jpg"]):
                structure["personalInfo"]["portfolio"] = url.strip()
                break

    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        structure["personalInfo"]["name"] = lines[0]

    # --- Section detection keywords ---
    SUMMARY_KEYS = ["summary", "profile", "objective", "professional summary", "career objective", "about me", "career summary"]
    SKILLS_KEYS = ["skills", "competencies", "technologies", "key skills", "core competencies", "expertise", "skills & tools", "technical skills", "skills and tools"]
    EXPERIENCE_KEYS = ["experience", "employment", "work history", "professional experience", "employment history", "career history", "work experience", "key responsibilities", "responsibilities", "internship", "internships"]
    PROJECTS_KEYS = ["projects", "accomplishments", "key projects", "technical projects", "academic projects", "personal projects"]
    EDUCATION_KEYS = ["education", "academic", "qualifications", "academic background", "credentials", "academic qualifications", "educational qualifications"]
    CERTIFICATIONS_KEYS = ["certifications", "licenses", "courses", "certification", "training", "trainings"]
    # Sections to skip entirely (content not added to any field)
    SKIP_KEYS = ["hobbies", "interests", "hobbies and interests", "personal profile", "personal details", "personal information", "declaration", "references", "extra curricular", "extracurricular", "co-curricular"]

    def detect_section(line_text: str) -> str:
        """Detects which section a header line belongs to. Returns section key or empty string."""
        # Strip common punctuation from headers
        cleaned = re.sub(r'[:\-–—|•*#_]+$', '', line_text).strip()
        lower = cleaned.lower()

        # Check each section type
        for key in SKIP_KEYS:
            if key in lower:
                return "skip"
        for key in SUMMARY_KEYS:
            if key in lower:
                return "summary"
        for key in SKILLS_KEYS:
            if key in lower:
                return "skills"
        for key in EXPERIENCE_KEYS:
            if key in lower:
                return "experience"
        for key in PROJECTS_KEYS:
            if key in lower:
                return "projects"
        for key in EDUCATION_KEYS:
            if key in lower:
                return "education"
        for key in CERTIFICATIONS_KEYS:
            if key in lower:
                return "certifications"
        return ""

    def is_section_header(line_text: str) -> bool:
        """Checks if a line looks like a section header."""
        stripped = line_text.strip()
        # All-caps or title-case short lines
        if len(stripped) < 60 and (stripped.isupper() or stripped.istitle() or stripped.endswith(":")):
            return True
        # Lines that are bold-style markers (all caps with possible colon/dash)
        cleaned = re.sub(r'[:\-–—|•*#_\s]+$', '', stripped)
        if len(cleaned) < 50 and cleaned.replace(" ", "").isalpha() and cleaned.isupper():
            return True
        return False

    current_section = ""
    temp_bullets: list = []
    temp_company = ""
    temp_role = ""
    temp_duration = ""

    def save_experience():
        nonlocal temp_role, temp_company, temp_duration, temp_bullets
        if temp_role or temp_company or temp_bullets:
            structure["experience"].append({
                "role": temp_role or "Specialist",
                "company": temp_company or "",
                "duration": temp_duration or "",
                "bullets": temp_bullets if temp_bullets else []
            })
            temp_role, temp_company, temp_duration, temp_bullets = "", "", "", []

    for line in lines[1:]:
        lower_line = line.lower()

        # --- Detect section header ---
        detected = detect_section(line)
        if detected and is_section_header(line):
            # Save any pending experience before switching
            if current_section == "experience":
                save_experience()
            current_section = detected
            continue

        # Also detect headers that aren't perfectly formatted but match keywords
        if detected and len(line) < 60:
            if current_section == "experience":
                save_experience()
            current_section = detected
            continue

        # --- Skip section: don't store content ---
        if current_section == "skip":
            continue

        # --- Parse content based on current section ---
        BULLET_CHARS = ["•", "●", "▪", "■", "○", "⁃", "–", "—", "*", "-", "✓", "✔"]
        BULLET_PATTERN = r'^[' + ''.join(re.escape(c) for c in BULLET_CHARS) + r'\s]+'
        
        # Check if line starts with a bullet character
        starts_with_bullet = any(line.startswith(c) for c in BULLET_CHARS)

        if current_section == "summary":
            structure["summary"] += (" " if structure["summary"] else "") + line

        elif current_section == "skills":
            # Split by common skill delimiters and strip any prepended bullets
            skills = [re.sub(BULLET_PATTERN, '', s).strip() for s in re.split(r'[,|•●\t]', line) if s.strip()]
            structure["skills"].extend([s for s in skills if s])

        elif current_section == "experience":
            # Check if line contains a year range (likely a new role header)
            if re.search(r'(19|20)\d{2}', line) or "present" in lower_line:
                save_experience()
                duration_match = re.search(r'((19|20)\d{2}\s*(-\s*(present|(19|20)\d{2}))?)', line, re.IGNORECASE)
                temp_duration = duration_match.group(0) if duration_match else ""
                cleaned_line = line.replace(temp_duration, "").strip() if temp_duration else line
                cleaned_line = re.sub(r'^[\-–—|•●*\s]+|[\-–—|•●*\s]+$', '', cleaned_line).strip()
                headers = [h.strip() for h in re.split(r'[,|•●\t]', cleaned_line) if h.strip()]
                temp_role = headers[0] if len(headers) > 0 else ""
                temp_company = headers[1] if len(headers) > 1 else ""
            elif starts_with_bullet:
                temp_bullets.append(re.sub(BULLET_PATTERN, '', line))
            else:
                # Continuation of a bullet or a standalone responsibility line
                if temp_bullets:
                    temp_bullets[-1] += " " + line
                else:
                    temp_bullets.append(line)

        elif current_section == "projects":
            if starts_with_bullet:
                bullet_text = re.sub(BULLET_PATTERN, '', line)
                if structure["projects"]:
                    structure["projects"][-1]["bullets"].append(bullet_text)
                else:
                    # No project header yet, create one
                    structure["projects"].append({
                        "name": "Project",
                        "description": "",
                        "bullets": [bullet_text]
                    })
            else:
                # Check if this is a new project name (contains | or is a short non-bullet line)
                is_new_project = ("|" in line) or (not structure["projects"]) or (
                    structure["projects"] and not structure["projects"][-1]["bullets"]
                )
                if is_new_project:
                    parts = [h.strip() for h in re.split(r'[|]', line) if h.strip()]
                    structure["projects"].append({
                        "name": parts[0] if len(parts) > 0 else "Project",
                        "description": parts[1] if len(parts) > 1 else "",
                        "bullets": []
                    })
                else:
                    if structure["projects"] and structure["projects"][-1]["bullets"]:
                        structure["projects"][-1]["bullets"][-1] += " " + line

        elif current_section == "education":
            # --- Handle tabular education data ---
            # Detect table header rows and skip them
            table_header_keys = ["examination", "board/university", "school/college", "year", "score", "grade", "cgpa", "percentage"]
            if sum(1 for k in table_header_keys if k in lower_line) >= 2:
                continue  # Skip table header row

            # Try to detect structured education entry with year
            edu_keywords = [
                "b.tech", "b.e.", "m.tech", "b.s.", "m.s.", "b.sc", "m.sc",
                "bachelor", "master", "ph.d", "degree", "diploma",
                "12th", "10th", "hsc", "ssc", "cssc", "iti",
                "high school", "school", "intermediate", "matriculation",
                "pass", "science", "commerce", "arts", "engineering"
            ]
            is_new_edu = any(k in lower_line for k in edu_keywords)
            year_match = re.search(r'\b((19|20)\d{2}\s*(-\s*(present|(19|20)\d{2}))?)\b', line, re.IGNORECASE)
            year_val = year_match.group(0) if year_match else ""

            # Try to extract score/grade/percentage/CGPA
            score_match = re.search(r'(?:CGPA|GPA|Percentage|Score|Grade)\s*[:\-]?\s*([\d.]+%?)', line, re.IGNORECASE)
            score_val = ""
            if score_match:
                score_val = score_match.group(0).strip()
            else:
                # Look for standalone percentage or CGPA pattern
                pct_match = re.search(r'\b(\d{1,3}(?:\.\d+)?)\s*%', line)
                cgpa_match = re.search(r'\b(\d\.\d{1,2})\b', line)
                if pct_match:
                    score_val = f"Percentage: {pct_match.group(0)}"
                elif cgpa_match and float(cgpa_match.group(1)) <= 10:
                    score_val = f"CGPA: {cgpa_match.group(1)}"

            # Clean the line by removing year and score
            cleaned_line = line
            if year_val:
                cleaned_line = cleaned_line.replace(year_val, "").strip()
            if score_val:
                # Remove score portion from display line
                cleaned_line = re.sub(re.escape(score_match.group(0) if score_match else (pct_match.group(0) if pct_match else "")), '', cleaned_line).strip() if (score_match or pct_match) else cleaned_line
            cleaned_line = re.sub(r'^[\-–—|•●*\s,]+|[\-–—|•●*\s,]+$', '', cleaned_line).strip()

            if is_new_edu and cleaned_line:
                # Split degree vs school using school keywords
                school_keywords = ["university", "college", "school", "institute", "bse", "cbse", "icse"]
                school_idx = -1
                lower_clean = cleaned_line.lower()
                for skw in school_keywords:
                    pos = lower_clean.find(skw)
                    if pos != -1 and (school_idx == -1 or pos < school_idx):
                        school_idx = pos
                
                if school_idx != -1:
                    # Capture the word before the school keyword if capitalized
                    pre_words = cleaned_line[:school_idx].strip().split()
                    exclude_prev = ["pass", "class", "degree", "diploma", "bachelor", "master", "doctor", "-"]
                    if pre_words and pre_words[-1][0].isupper() and pre_words[-1].lower() not in exclude_prev:
                        school_part = pre_words[-1]
                        degree_text = " ".join(pre_words[:-1]).strip()
                        school_text = school_part + " " + cleaned_line[school_idx:].strip()
                    else:
                        degree_text = cleaned_line[:school_idx].strip()
                        school_text = cleaned_line[school_idx:].strip()
                else:
                    parts = [p.strip() for p in re.split(r'[,\t|–—]', cleaned_line) if p.strip()]
                    degree_text = parts[0] if parts else cleaned_line
                    school_text = " — ".join(parts[1:]) if len(parts) > 1 else ""

                year_display = year_val
                if score_val:
                    year_display = f"{year_val}  |  {score_val}" if year_val else score_val

                structure["education"].append({
                    "degree": degree_text,
                    "school": school_text,
                    "year": year_display
                })
            else:
                # Continuation line — append to the last education entry
                if structure["education"] and cleaned_line:
                    last_edu = structure["education"][-1]
                    if not last_edu["school"]:
                        last_edu["school"] = cleaned_line
                    else:
                        last_edu["school"] += " — " + cleaned_line
                    if year_val and not last_edu["year"]:
                        last_edu["year"] = year_val
                    if score_val and score_val not in last_edu["year"]:
                        last_edu["year"] = f"{last_edu['year']}  |  {score_val}" if last_edu["year"] else score_val
                elif cleaned_line:
                    structure["education"].append({
                        "degree": cleaned_line,
                        "school": "",
                        "year": year_val or ""
                    })

        elif current_section == "certifications":
            cert_val = re.sub(BULLET_PATTERN, '', line)
            if cert_val:
                structure["certifications"].append(cert_val)


    # Save any remaining experience
    save_experience()

    # Default fallbacks
    if not structure["summary"]:
        structure["summary"] = "Dedicated specialist seeking target growth role."
    if not structure["skills"]:
        structure["skills"] = ["Problem Solving", "Collaboration", "Communication"]

    return structure
