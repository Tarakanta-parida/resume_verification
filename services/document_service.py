import io
import re
import copy
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any, List

def clean_html_tags(text: str) -> str:
    if not text:
        return ""
    return re.sub(r'</?[a-zA-Z]+(?:\s+[^>]*)?>', '', text)

def clean_contact_value(val: str) -> str:
    if not val:
        return ""
    # Strip leading/trailing spaces, dashes, colons, vertical bars, bullets
    cleaned = re.sub(r'^[\s\-:|•·▪◦-]+|[\s\-:|•·▪◦-]+$', '', val)
    cleaned = cleaned.strip()
    # Hide default placeholders
    lower_val = cleaned.lower()
    if "linkedin.com/in/user" in lower_val or "github.com/user" in lower_val or "info@domain.com" in lower_val:
        return ""
    return cleaned


def process_bullet_text(text: str, show_added: bool, show_optimized: bool, revert_entire_string: bool = False) -> str:
    if not text:
        return text

    result = text

    # 1. Process optimized bullets (class="mod")
    if not show_optimized:
        if revert_entire_string:
            match = re.search(r'<mark\s+class="mod"\s+data-tooltip="Original:\s*([^"]+)"[^>]*>', result, flags=re.IGNORECASE)
            if not match:
                match = re.search(r'<mark\s+class="mod"\s+data-tooltip=\'Original:\s*([^\']+)\'[^>]*>', result, flags=re.IGNORECASE)
            if match:
                return match.group(1)

        result = re.sub(r'<mark\s+class="mod"\s+data-tooltip="Original:\s*([^"]+)"[^>]*>.*?</mark>', r'\1', result, flags=re.IGNORECASE)
        result = re.sub(r'<mark\s+class="mod"\s+data-tooltip=\'Original:\s*([^\']+)\'[^>]*>.*?</mark>', r'\1', result, flags=re.IGNORECASE)

    # 2. Process added keywords (class="add")
    if not show_added:
        result = re.sub(r'<mark\s+class="add"[^>]*>.*?</mark>', '', result, flags=re.IGNORECASE)
        result = re.sub(r'\s+and\s*([.,;!])', r'\1', result, flags=re.IGNORECASE)
        result = re.sub(r'\s+or\s*([.,;!])', r'\1', result, flags=re.IGNORECASE)
        result = re.sub(r',\s*and\s*([.,;!])', r'\1', result, flags=re.IGNORECASE)
        result = re.sub(r',\s*or\s*([.,;!])', r'\1', result, flags=re.IGNORECASE)
        result = re.sub(r'\s+,\s*', ', ', result)
        result = re.sub(r'\s+', ' ', result)
        result = re.sub(r'\s+([.,;!])', r'\1', result)
        result = re.sub(r'\s*•\s*•\s*', ' • ', result)
        result = result.strip()

    return result


def process_resume_data_highlights(data: dict, original_data: dict, show_added: bool, show_optimized: bool) -> dict:
    processed = copy.deepcopy(data)

    if "summary" in processed:
        summary_text = original_data.get("summary", "") if (not show_optimized and original_data) else processed["summary"]
        processed["summary"] = process_bullet_text(summary_text, show_added, show_optimized)

    if "skills" in processed and isinstance(processed["skills"], list):
        new_skills = []
        for s in processed["skills"]:
            p_s = process_bullet_text(s, show_added, show_optimized)
            if p_s.strip():
                new_skills.append(p_s)
        processed["skills"] = new_skills

    if "experience" in processed and isinstance(processed["experience"], list):
        orig_exp_list = original_data.get("experience", []) if original_data else []
        for i, exp in enumerate(processed["experience"]):
            if "role" in exp:
                exp["role"] = process_bullet_text(exp["role"], show_added, show_optimized)
            if "bullets" in exp and isinstance(exp["bullets"], list):
                new_bullets = []
                orig_bullets = orig_exp_list[i].get("bullets", []) if (i < len(orig_exp_list) and isinstance(orig_exp_list[i], dict)) else []
                for bIdx, b in enumerate(exp["bullets"]):
                    bullet_text = orig_bullets[bIdx] if (not show_optimized and bIdx < len(orig_bullets)) else b
                    new_bullets.append(process_bullet_text(bullet_text, show_added, show_optimized, revert_entire_string=True))
                exp["bullets"] = new_bullets

    if "projects" in processed and isinstance(processed["projects"], list):
        orig_proj_list = original_data.get("projects", []) if original_data else []
        for i, proj in enumerate(processed["projects"]):
            if "description" in proj:
                proj["description"] = process_bullet_text(proj["description"], show_added, show_optimized)
            if "bullets" in proj and isinstance(proj["bullets"], list):
                new_bullets = []
                orig_bullets = orig_proj_list[i].get("bullets", []) if (i < len(orig_proj_list) and isinstance(orig_proj_list[i], dict)) else []
                for bIdx, b in enumerate(proj["bullets"]):
                    bullet_text = orig_bullets[bIdx] if (not show_optimized and bIdx < len(orig_bullets)) else b
                    new_bullets.append(process_bullet_text(bullet_text, show_added, show_optimized, revert_entire_string=True))
                proj["bullets"] = new_bullets

    if "certifications" in processed and isinstance(processed["certifications"], list):
        new_certs = []
        orig_certs = original_data.get("certifications", []) if original_data else []
        for i, cert in enumerate(processed["certifications"]):
            cert_text = orig_certs[i] if (not show_optimized and i < len(orig_certs)) else cert
            new_certs.append(process_bullet_text(cert_text, show_added, show_optimized))
        processed["certifications"] = new_certs

    return processed


def generate_styled_resume_html(data: Dict[str, Any], include_highlights: bool = False) -> str:
    """Renders resume data structure into clean, single-page print-ready HTML."""
    skills_list = []
    for s in data.get("skills", []):
        if not include_highlights:
            s = clean_html_tags(s)
        skills_list.append(s)
    skills_html = " • ".join(skills_list)

    exp_html_runs = []
    for exp in data.get("experience", []):
        bullets_html = ""
        for b in exp.get("bullets", []):
            if not include_highlights:
                b = clean_html_tags(b)
            bullets_html += f"<li>{b}</li>"
        role = exp.get("role", "")
        if not include_highlights:
            role = clean_html_tags(role)
        exp_html_runs.append(f"""
        <div class="experience-item">
            <div class="item-header">
                <span class="bold-text">{role}</span>
                <span>{exp.get("duration", "")}</span>
            </div>
            <div class="item-sub">
                <span>{exp.get("company", "")}</span>
            </div>
            <ul>{bullets_html}</ul>
        </div>
        """)
    exp_html = "\n".join(exp_html_runs)

    proj_html_runs = []
    for proj in data.get("projects", []):
        bullets_html = ""
        for b in proj.get("bullets", []):
            if not include_highlights:
                b = clean_html_tags(b)
            bullets_html += f"<li>{b}</li>"
        proj_html_runs.append(f"""
        <div class="project-item">
            <div class="item-header">
                <span class="bold-text">{proj.get("name", "")}</span>
            </div>
            <div class="item-sub">
                <span>{proj.get("description", "")}</span>
            </div>
            <ul>{bullets_html}</ul>
        </div>
        """)
    proj_html = "\n".join(proj_html_runs)

    edu_html_runs = []
    for edu in data.get("education", []):
        edu_html_runs.append(f"""
        <div class="experience-item">
            <div class="item-header">
                <span class="bold-text">{edu.get("degree", "")}</span>
                <span>{edu.get("year", "")}</span>
            </div>
            <div class="item-sub">
                <span>{edu.get("school", "")}</span>
            </div>
        </div>
        """)
    edu_html = "\n".join(edu_html_runs)

    cert_html_runs = []
    for cert in data.get("certifications", []):
        if not include_highlights:
            cert = clean_html_tags(cert)
        cert_html_runs.append(f"<li>{cert}</li>")
    cert_html = "\n".join(cert_html_runs)

    personal = data.get("personalInfo", {})
    name = personal.get("name", "Applicant Name")
    email = clean_contact_value(personal.get("email", ""))
    phone = clean_contact_value(personal.get("phone", ""))
    linkedin = clean_contact_value(personal.get("linkedin", ""))
    github = clean_contact_value(personal.get("github", ""))
    portfolio = clean_contact_value(personal.get("portfolio", ""))

    contacts = []
    if email: contacts.append(email)
    if phone: contacts.append(phone)
    if linkedin: contacts.append(linkedin)
    if github: contacts.append(github)
    if portfolio: contacts.append(portfolio)
    contacts_str = " | ".join(contacts)

    summary = data.get("summary", "")
    if not include_highlights:
        summary = clean_html_tags(summary)

    return f"""
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <title>{name} - Optimized Resume</title>
        <style>
            body {{
                font-family: 'Arial', sans-serif;
                line-height: 1.5;
                color: #2D3748;
                margin: 0;
                padding: 40px;
                background-color: #FFFFFF;
                font-size: 11pt;
            }}
            h2 {{
                font-size: 18pt;
                font-weight: bold;
                text-align: center;
                color: #1A202C;
                margin-top: 0;
                margin-bottom: 4px;
            }}
            .personal-info {{
                text-align: center;
                font-size: 9pt;
                color: #718096;
                margin-bottom: 20px;
                border-bottom: 1px solid #E2E8F0;
                padding-bottom: 8px;
            }}
            h3 {{
                font-size: 11pt;
                font-weight: bold;
                text-transform: uppercase;
                border-bottom: 1px solid #A0AEC0;
                padding-bottom: 3px;
                margin-top: 18px;
                margin-bottom: 8px;
                color: #1A202C;
            }}
            .section-summary p {{
                color: #4A5568;
                text-align: justify;
                margin: 0;
            }}
            .skills-list {{
                color: #4A5568;
            }}
            .experience-item, .project-item {{
                margin-bottom: 12px;
            }}
            .item-header {{
                display: flex;
                justify-content: space-between;
                font-weight: bold;
                color: #1A202C;
                font-size: 11pt;
            }}
            .item-sub {{
                display: flex;
                justify-content: space-between;
                color: #718096;
                font-style: italic;
                font-size: 9pt;
                margin-bottom: 4px;
            }}
            .bold-text {{
                font-weight: bold;
            }}
            ul {{
                margin-top: 2px;
                margin-bottom: 0;
                padding-left: 20px;
                color: #4A5568;
            }}
            li {{
                margin-bottom: 3px;
            }}
            mark.add {{
                background-color: rgba(16, 185, 129, 0.2);
                color: #047857;
                border-radius: 2px;
            }}
            mark.mod {{
                background-color: rgba(245, 158, 11, 0.2);
                color: #B45309;
                border-radius: 2px;
            }}
            @media print {{
                body {{
                    padding: 0;
                }}
            }}
        </style>
    </head>
    <body>
        <h2>{name}</h2>
        <div class="personal-info">
            {contacts_str}
        </div>
        
        <h3>Professional Summary</h3>
        <div class="section-summary">
            <p>{summary}</p>
        </div>
        
        <h3>Technical Skills</h3>
        <div class="skills-list">
            {skills_html}
        </div>
        
        <h3>Work Experience</h3>
        {exp_html}
        
        {f"<h3>Technical Projects</h3>{proj_html}" if proj_html.strip() else ""}
        
        <h3>Education</h3>
        {edu_html}
        
        {f"<h3>Certifications</h3><ul>{cert_html}</ul>" if cert_html.strip() else ""}
    </body>
    </html>
    """


def generate_docx_from_data(data: dict) -> io.BytesIO:
    """Generates clean formatted DOCX byte stream from resume JSON."""
    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'

    def add_body_paragraph(text, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        return p

    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_after = Pt(2)
    run_name = p_header.add_run(data.get("personalInfo", {}).get("name", "Name").upper())
    run_name.bold = True
    run_name.font.size = Pt(16)
    run_name.font.name = 'Arial'

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(12)
    contact_parts = []
    personal = data.get("personalInfo", {})
    email = clean_contact_value(personal.get("email", ""))
    phone = clean_contact_value(personal.get("phone", ""))
    linkedin = clean_contact_value(personal.get("linkedin", ""))
    github = clean_contact_value(personal.get("github", ""))
    portfolio = clean_contact_value(personal.get("portfolio", ""))

    if email: contact_parts.append(email)
    if phone: contact_parts.append(phone)
    if linkedin: contact_parts.append(linkedin)
    if github: contact_parts.append(github)
    if portfolio: contact_parts.append(portfolio)

    run_contact = p_contact.add_run("  |  ".join(contact_parts))
    run_contact.font.size = Pt(9.5)
    run_contact.font.name = 'Arial'

    if data.get("summary"):
        add_section_heading("Professional Summary")
        clean_summary = clean_html_tags(data["summary"])
        add_body_paragraph(clean_summary, space_after=12)

    if data.get("skills"):
        add_section_heading("Technical Skills")
        clean_skills = [clean_html_tags(s) for s in data["skills"]]
        skills_text = "  •  ".join(clean_skills)
        add_body_paragraph(skills_text, space_after=12)

    if data.get("experience"):
        add_section_heading("Work Experience")
        for exp in data["experience"]:
            role_clean = clean_html_tags(exp.get("role", ""))
            p_exp_header = doc.add_paragraph()
            p_exp_header.paragraph_format.space_after = Pt(1)
            p_exp_header.paragraph_format.keep_with_next = True

            run_role = p_exp_header.add_run(f"{role_clean}  |  ")
            run_role.bold = True
            run_role.font.name = 'Arial'
            run_role.font.size = Pt(11)

            run_company = p_exp_header.add_run(f"{exp.get('company', '')}")
            run_company.font.italic = True
            run_company.font.name = 'Arial'
            run_company.font.size = Pt(10.5)

            run_duration = p_exp_header.add_run(f" ({exp.get('duration', '')})")
            run_duration.font.name = 'Arial'
            run_duration.font.size = Pt(10)

            for bullet in exp.get("bullets", []):
                bullet_clean = clean_html_tags(bullet)
                p_bullet = doc.add_paragraph(style='List Bullet')
                p_bullet.paragraph_format.space_after = Pt(2)
                p_bullet.paragraph_format.line_spacing = 1.15
                run_bullet = p_bullet.add_run(bullet_clean)
                run_bullet.font.name = 'Arial'
                run_bullet.font.size = Pt(10.5)

    if data.get("projects"):
        non_empty_projects = [p for p in data["projects"] if p.get("name")]
        if non_empty_projects:
            add_section_heading("Technical Projects")
            for proj in non_empty_projects:
                desc_clean = clean_html_tags(proj.get("description", ""))
                p_proj_header = doc.add_paragraph()
                p_proj_header.paragraph_format.space_after = Pt(1)
                p_proj_header.paragraph_format.keep_with_next = True

                run_proj_name = p_proj_header.add_run(proj.get("name", ""))
                run_proj_name.bold = True
                run_proj_name.font.name = 'Arial'
                run_proj_name.font.size = Pt(11)

                if desc_clean:
                    run_proj_desc = p_proj_header.add_run(f"  |  {desc_clean}")
                    run_proj_desc.font.italic = True
                    run_proj_desc.font.name = 'Arial'
                    run_proj_desc.font.size = Pt(10.5)

                for bullet in proj.get("bullets", []):
                    bullet_clean = clean_html_tags(bullet)
                    p_bullet = doc.add_paragraph(style='List Bullet')
                    p_bullet.paragraph_format.space_after = Pt(2)
                    p_bullet.paragraph_format.line_spacing = 1.15
                    run_bullet = p_bullet.add_run(bullet_clean)
                    run_bullet.font.name = 'Arial'
                    run_bullet.font.size = Pt(10.5)

    if data.get("education"):
        add_section_heading("Education")
        for edu in data["education"]:
            p_edu = doc.add_paragraph()
            p_edu.paragraph_format.space_after = Pt(3)

            run_degree = p_edu.add_run(f"{edu.get('degree', '')}  |  ")
            run_degree.bold = True
            run_degree.font.name = 'Arial'
            run_degree.font.size = Pt(11)

            run_school = p_edu.add_run(f"{edu.get('school', '')}")
            run_school.font.italic = True
            run_school.font.name = 'Arial'
            run_school.font.size = Pt(10.5)

            run_year = p_edu.add_run(f" ({edu.get('year', '')})")
            run_year.font.name = 'Arial'
            run_year.font.size = Pt(10)

    if data.get("certifications"):
        add_section_heading("Certifications")
        for cert in data["certifications"]:
            cert_clean = clean_html_tags(cert)
            p_cert = doc.add_paragraph(style='List Bullet')
            p_cert.paragraph_format.space_after = Pt(2)
            p_cert.paragraph_format.line_spacing = 1.15
            run_cert = p_cert.add_run(cert_clean)
            run_cert.font.name = 'Arial'
            run_cert.font.size = Pt(10.5)

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def generate_plain_text(data: dict, original_data: dict | None = None, show_added: bool = True, show_optimized: bool = True) -> str:
    """Generates formatted plain text version of the resume."""
    personal = data.get("personalInfo", {})
    name = personal.get("name", "Applicant Name").upper()
    email = clean_contact_value(personal.get("email", ""))
    phone = clean_contact_value(personal.get("phone", ""))
    linkedin = clean_contact_value(personal.get("linkedin", ""))
    github = clean_contact_value(personal.get("github", ""))
    portfolio = clean_contact_value(personal.get("portfolio", ""))

    contact_parts = []
    if email: contact_parts.append(f"Email: {email}")
    if phone: contact_parts.append(f"Phone: {phone}")
    if linkedin: contact_parts.append(f"LinkedIn: {linkedin}")
    if github: contact_parts.append(f"GitHub: {github}")
    if portfolio: contact_parts.append(f"Portfolio: {portfolio}")
    contacts_str = " | ".join(contact_parts)

    summary_text = original_data.get("summary", "") if (not show_optimized and original_data) else data.get("summary", "")
    clean_summary = clean_html_tags(process_bullet_text(str(summary_text or ""), show_added, show_optimized))

    skills = [clean_html_tags(process_bullet_text(s, show_added, show_optimized)) for s in data.get("skills", [])]
    skills_clean = [s for s in skills if s.strip()]
    skills_str = " • ".join(skills_clean)

    exp_text = ""
    for i, exp in enumerate(data.get("experience", [])):
        role = clean_html_tags(process_bullet_text(exp.get("role", ""), show_added, show_optimized))
        company = exp.get("company", "")
        duration = exp.get("duration", "")

        bullets_str = ""
        orig_bullets = original_data.get("experience", [])[i].get("bullets", []) if (original_data and i < len(original_data.get("experience", []))) else []
        for bIdx, b in enumerate(exp.get("bullets", [])):
            bullet_text = orig_bullets[bIdx] if (not show_optimized and bIdx < len(orig_bullets)) else b
            b_clean = clean_html_tags(process_bullet_text(bullet_text, show_added, show_optimized, revert_entire_string=True))
            bullets_str += f"* {b_clean}\n"

        exp_text += f"\n{role} | {company} ({duration})\n{bullets_str}"

    proj_text = ""
    for i, proj in enumerate(data.get("projects", [])):
        p_name = proj.get("name", "")
        p_desc = clean_html_tags(process_bullet_text(proj.get("description", ""), show_added, show_optimized))

        bullets_str = ""
        orig_bullets = original_data.get("projects", [])[i].get("bullets", []) if (original_data and i < len(original_data.get("projects", []))) else []
        for bIdx, b in enumerate(proj.get("bullets", [])):
            bullet_text = orig_bullets[bIdx] if (not show_optimized and bIdx < len(orig_bullets)) else b
            b_clean = clean_html_tags(process_bullet_text(bullet_text, show_added, show_optimized, revert_entire_string=True))
            bullets_str += f"* {b_clean}\n"

        proj_text += f"\n{p_name} | {p_desc}\n{bullets_str}"

    edu_text = ""
    for edu in data.get("education", []):
        edu_text += f"\n{edu.get('degree', '')} | {edu.get('school', '')} ({edu.get('year', '')})\n"

    cert_text = ""
    orig_certs = original_data.get("certifications", []) if original_data else []
    for i, cert in enumerate(data.get("certifications", [])):
        c_val = orig_certs[i] if (not show_optimized and i < len(orig_certs)) else cert
        c_clean = clean_html_tags(process_bullet_text(c_val, show_added, show_optimized))
        cert_text += f"* {c_clean}\n"

    return f"""
{name}
{email} | {phone}

PROFESSIONAL SUMMARY
-------------------------------------
{clean_summary}

TECHNICAL SKILLS
-------------------------------------
{skills_str}

WORK EXPERIENCE
-------------------------------------
{exp_text}
{f"TECHNICAL PROJECTS\n-------------------------------------\n{proj_text}" if proj_text.strip() else ""}
EDUCATION
-------------------------------------
{edu_text}
{f"CERTIFICATIONS\n-------------------------------------\n{cert_text}" if cert_text.strip() else ""}
""".strip()
